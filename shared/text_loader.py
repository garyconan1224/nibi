from __future__ import annotations

"""文本输入层加载器（Phase 2C.1）。

统一封装 PDF / DOCX / 网页正文 三类素材的抽取，返回归一的 ``TextDocument``。
- PDF  → pypdf.PdfReader
- DOCX → python-docx.Document
- URL  → httpx 拉取 + readability-lxml 抽正文

设计约束：
- 所有加载器把底层异常归一化为 ``TextLoaderError``，pipeline 层只需 catch 这一类。
- 不做 OCR、不做扫描件处理（pypdf 抽不到文字直接返回空 content + warning）。
- 不持久化，调用方自行落盘。
"""

from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Dict, Literal, Union

import httpx

SourceType = Literal["pdf", "docx", "url"]

# 默认 UA：用常见桌面浏览器，避免被站点直接 403；不绕过 robots / 不伪造来源
_DEFAULT_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0 Safari/537.36"
)
_DEFAULT_TIMEOUT_S = 30.0
_MAX_HTML_BYTES = 5 * 1024 * 1024  # 5MB 单页正文上限


class TextLoaderError(RuntimeError):
    """统一错误类型。原始异常通过 ``raise ... from err`` 链式保留。"""


@dataclass
class TextDocument:
    title: str
    content: str
    source_type: SourceType
    source: str                   # 原始 URL 或文件路径
    char_count: int
    meta: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


def _normalize(text: str) -> str:
    """去多余空白：保留段落结构（\n\n），但折叠行内多空格。"""
    if not text:
        return ""
    lines = [ln.rstrip() for ln in text.splitlines()]
    # 折叠 ≥3 个连续空行为 2 个
    out: list[str] = []
    blank_run = 0
    for ln in lines:
        if ln.strip() == "":
            blank_run += 1
            if blank_run <= 1:
                out.append("")
        else:
            blank_run = 0
            out.append(ln)
    return "\n".join(out).strip()


# ── PDF ───────────────────────────────────────────────────

# marker 模型加载（懒加载，进程内单例）
_marker_converter = None


def _get_marker_converter():
    """懒加载 marker PdfConverter（模型 ~1.5GB，首次约 30s）。"""
    global _marker_converter
    if _marker_converter is not None:
        return _marker_converter
    try:
        from marker.converters.pdf import PdfConverter
        from marker.config.parser import ConfigParser
        from marker.models import create_model_dict

        model_dict = create_model_dict()
        config_parser = ConfigParser({
            "output_format": "markdown",
            "disable_image_extraction": True,  # 图片保留占位符，不提取为文件
        })
        _marker_converter = PdfConverter(
            artifact_dict=model_dict,
            processor_list=config_parser.get_processors(),
            renderer=config_parser.get_renderer(),
        )
    except Exception as err:  # noqa: BLE001
        raise TextLoaderError(f"marker 初始化失败: {err}") from err
    return _marker_converter


def _load_pdf_marker(path: Path) -> TextDocument:
    """用 marker 解析 PDF → Markdown（保留图片+表格，支持扫描件 OCR）。"""
    converter = _get_marker_converter()
    try:
        result = converter(str(path))
    except Exception as err:  # noqa: BLE001
        raise TextLoaderError(f"marker PDF 解析失败: {err}") from err

    md_text = getattr(result, "markdown", "") or ""
    content = _normalize(md_text)
    title = _extract_pdf_title_pypdf(path) or path.stem

    return TextDocument(
        title=title or path.stem,
        content=content,
        source_type="pdf",
        source=str(path.resolve()),
        char_count=len(content),
        meta={"parser": "marker"},
    )


def _extract_pdf_title_pypdf(path: Path) -> str:
    """用 pypdf 提取 PDF 标题（轻量，不加载 marker 模型）。"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        meta_obj = reader.metadata or {}
        return str(meta_obj.get("/Title") or "").strip()
    except Exception:  # noqa: BLE001
        return ""


def load_pdf(path: Union[str, Path]) -> TextDocument:
    """解析 PDF 文件。marker 优先（支持扫描件 OCR + 图片表格保留），pypdf 兜底。"""
    p = Path(path)
    if not p.is_file():
        raise TextLoaderError(f"PDF 文件不存在: {p}")

    # 优先 marker
    try:
        return _load_pdf_marker(p)
    except TextLoaderError:
        pass  # marker 失败，fallback pypdf
    except Exception:  # noqa: BLE001
        pass  # marker 导入失败等

    # fallback: pypdf（纯文本，无 OCR）
    try:
        from pypdf import PdfReader
    except ImportError as err:
        raise TextLoaderError("缺少依赖 pypdf；请 pip install pypdf") from err

    try:
        reader = PdfReader(str(p))
    except Exception as err:  # noqa: BLE001
        raise TextLoaderError(f"PDF 解析失败: {err}") from err

    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as err:  # noqa: BLE001
            raise TextLoaderError(f"PDF 第 {i + 1} 页文本抽取失败: {err}") from err

    content = _normalize("\n\n".join(pages))
    title = ""
    try:
        meta_obj = reader.metadata or {}
        title = str(meta_obj.get("/Title") or "").strip()
    except Exception:  # noqa: BLE001
        title = ""
    if not title:
        title = p.stem

    return TextDocument(
        title=title,
        content=content,
        source_type="pdf",
        source=str(p.resolve()),
        char_count=len(content),
        meta={"parser": "pypdf", "page_count": len(reader.pages)},
    )


# ── DOCX ──────────────────────────────────────────────────


def load_docx(path: Union[str, Path]) -> TextDocument:
    p = Path(path)
    if not p.is_file():
        raise TextLoaderError(f"DOCX 文件不存在: {p}")
    try:
        from docx import Document  # python-docx
    except ImportError as err:
        raise TextLoaderError("缺少依赖 python-docx；请 pip install python-docx") from err

    try:
        doc = Document(str(p))
    except Exception as err:  # noqa: BLE001
        raise TextLoaderError(f"DOCX 解析失败: {err}") from err

    paragraphs = [para.text for para in doc.paragraphs]
    # 表格也按行拼接，便于 LLM 阅读
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    content = _normalize("\n\n".join(paragraphs))
    title = ""
    try:
        core = doc.core_properties
        title = (core.title or "").strip()
    except Exception:  # noqa: BLE001
        title = ""
    if not title:
        title = p.stem

    return TextDocument(
        title=title,
        content=content,
        source_type="docx",
        source=str(p.resolve()),
        char_count=len(content),
        meta={"paragraph_count": len(doc.paragraphs)},
    )


# ── URL（网页正文）────────────────────────────────────────


def load_url(url: str, *, timeout: float = _DEFAULT_TIMEOUT_S) -> TextDocument:
    if not url or not url.strip():
        raise TextLoaderError("URL 不能为空")
    url = url.strip()

    try:
        with httpx.Client(
            headers={"User-Agent": _DEFAULT_UA, "Accept": "text/html,*/*"},
            timeout=timeout,
            follow_redirects=True,
        ) as client:
            resp = client.get(url)
    except httpx.HTTPError as err:
        raise TextLoaderError(f"网页拉取失败: {err}") from err

    if resp.status_code >= 400:
        raise TextLoaderError(f"网页拉取失败: HTTP {resp.status_code}")

    raw = resp.content[:_MAX_HTML_BYTES]
    if not raw:
        raise TextLoaderError("网页返回空内容")

    try:
        from readability import Document as ReadabilityDoc
    except ImportError as err:
        raise TextLoaderError("缺少依赖 readability-lxml；请 pip install readability-lxml") from err

    # readability 期望字符串；httpx.Response.text 会按 charset / chardet 解码
    html_text = resp.text if isinstance(resp.text, str) else raw.decode("utf-8", errors="replace")

    try:
        doc = ReadabilityDoc(html_text)
        title = (doc.short_title() or doc.title() or "").strip()
        summary_html = doc.summary(html_partial=True) or ""
    except Exception as err:  # noqa: BLE001
        raise TextLoaderError(f"网页正文抽取失败: {err}") from err

    # 把 HTML 摘要节点转为纯文本（lxml 的 text_content 拿全部可见文字）
    try:
        from lxml import html as lxml_html  # type: ignore[import-not-found]
        tree = lxml_html.fromstring(summary_html) if summary_html else None
        text = tree.text_content() if tree is not None else ""
    except Exception as err:  # noqa: BLE001
        raise TextLoaderError(f"网页正文转文本失败: {err}") from err

    content = _normalize(text)
    if not title:
        title = url

    return TextDocument(
        title=title,
        content=content,
        source_type="url",
        source=url,
        char_count=len(content),
        meta={
            "http_status": resp.status_code,
            "final_url": str(resp.url),
            "content_type": resp.headers.get("content-type", ""),
        },
    )


# ── 自动分派 ──────────────────────────────────────────────


def load_auto(source: str, source_type: str | None = None) -> TextDocument:
    """根据 source_type / 文件扩展名自动选择 loader。

    优先级：显式 source_type > URL 前缀 > 扩展名。
    """
    explicit = (source_type or "").strip().lower()
    if explicit == "url":
        return load_url(source)
    if explicit == "pdf":
        return load_pdf(source)
    if explicit == "docx":
        return load_docx(source)

    s = (source or "").strip()
    if s.lower().startswith(("http://", "https://")):
        return load_url(s)

    ext = Path(s).suffix.lower()
    if ext == ".pdf":
        return load_pdf(s)
    if ext == ".docx":
        return load_docx(s)
    raise TextLoaderError(
        f"无法判断文本来源类型，请显式提供 source_type=pdf|docx|url（source={s!r}）"
    )
