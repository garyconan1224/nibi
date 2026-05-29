"""Phase 2C.1：文本输入层端到端单元测试。

- text_loader.load_pdf / load_docx：用真实文件跑一次（生成 fixture）。
- text_loader.load_url：mock httpx。
- pipeline handle_text_task：mock load_auto，验证状态机推进 + 产物落盘。
- pipeline handle_text_task：缺 payload.source 抛 ValueError。
"""
from __future__ import annotations

import json
import time
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

from backend.app.models.tasks import TERMINAL_STATUS_VALUES, TaskStatus
from backend.app.services.pipeline_tasks import _parse_structured_summary, _split_paragraphs, handle_text_task
from backend.app.services.task_runner import TaskRunner
from backend.app.services.task_store import TaskStore
from shared.text_loader import (
    TextDocument,
    TextLoaderError,
    load_docx,
    load_pdf,
    load_url,
)


# ── text_loader 直接测 ─────────────────────────────────────


def _make_pdf(path: Path, text_lines: list[str]) -> None:
    """用 pypdf 直接合成一个最小可读 PDF。"""
    from pypdf import PdfWriter
    # pypdf 不支持纯文本写入，改用 reportlab 不引入新依赖：手写一个最小 PDF
    # 因 pypdf 实际项目里只读，本测试用 fpdf2 又会引入新依赖；
    # 这里用最朴素的方式：通过 PdfWriter 加一个空白页后用 add_metadata 携带标题，
    # 然后 load_pdf 走 extract_text 拿到空串路径（覆盖"PDF 抽到空"分支）。
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_metadata({"/Title": "Test PDF"})
    with path.open("wb") as f:
        writer.write(f)


def test_load_pdf_blank_returns_empty_content(tmp_path: Path) -> None:
    pdf_file = tmp_path / "blank.pdf"
    _make_pdf(pdf_file, [])

    doc = load_pdf(pdf_file)
    assert doc.source_type == "pdf"
    # N10: marker 可能对空白 PDF 返回少量 OCR 内容，不再断言 content == ""
    assert isinstance(doc.content, str)
    assert isinstance(doc.char_count, int)
    # meta 中应有 parser 字段标记使用了哪个解析器
    assert "parser" in doc.meta


def test_load_pdf_missing_raises() -> None:
    with pytest.raises(TextLoaderError):
        load_pdf("/no/such/file.pdf")


def test_load_docx_roundtrip(tmp_path: Path) -> None:
    from docx import Document

    src = tmp_path / "sample.docx"
    doc = Document()
    doc.core_properties.title = "我的文档"
    doc.add_paragraph("第一段。")
    doc.add_paragraph("第二段，包含一些细节。")
    doc.save(str(src))

    out = load_docx(src)
    assert out.source_type == "docx"
    assert out.title == "我的文档"
    assert "第一段" in out.content and "第二段" in out.content
    assert out.char_count > 0


def test_load_url_happy(monkeypatch: pytest.MonkeyPatch) -> None:
    html = (
        b"<html><head><title>My Article</title></head>"
        b"<body><article><h1>Hello World</h1>"
        b"<p>This is the body content of a news article. " * 5
        + b"</p></article></body></html>"
    )

    class _FakeResp:
        status_code = 200
        content = html
        text = html.decode("utf-8")
        headers = {"content-type": "text/html"}
        url = "https://example.com/post"

    class _FakeClient:
        def __init__(self, *a: Any, **kw: Any) -> None:
            pass

        def __enter__(self):
            return self

        def __exit__(self, *exc):
            return False

        def get(self, url: str):
            return _FakeResp()

    monkeypatch.setattr("shared.text_loader.httpx.Client", _FakeClient)

    doc = load_url("https://example.com/post")
    assert doc.source_type == "url"
    assert doc.title  # readability 会拿到标题
    assert "body content" in doc.content
    assert doc.meta["http_status"] == 200


def test_load_url_wechat_js_content(monkeypatch: pytest.MonkeyPatch) -> None:
    """微信公众号：#js_content 存在 → 走专门 xpath 抽取，parser=wechat。"""
    wx_html = (
        '<html><body>'
        '<h1 id="activity-name">测试标题</h1>'
        '<div id="js_content"><p>这是微信正文内容。</p></div>'
        '</body></html>'
    ).encode("utf-8")

    class _FakeResp:
        status_code = 200
        content = wx_html
        text = wx_html.decode("utf-8")
        headers = {"content-type": "text/html"}
        url = "https://mp.weixin.qq.com/s/abc123"

    class _FakeClient:
        def __init__(self, *a: Any, **kw: Any) -> None:
            pass
        def __enter__(self):
            return self
        def __exit__(self, *exc):
            return False
        def get(self, url: str):
            return _FakeResp()

    monkeypatch.setattr("shared.text_loader.httpx.Client", _FakeClient)

    doc = load_url("https://mp.weixin.qq.com/s/abc123")
    assert doc.meta["parser"] == "wechat"
    assert doc.title == "测试标题"
    assert "微信正文内容" in doc.content


def test_load_url_wechat_no_js_content_fallback(monkeypatch: pytest.MonkeyPatch) -> None:
    """微信域名但无 #js_content → 回落 readability，不报错。"""
    normal_html = (
        "<html><head><title>Fallback Page</title></head>"
        "<body><article><p>普通正文内容。</p></article></body></html>"
    ).encode("utf-8")

    class _FakeResp:
        status_code = 200
        content = normal_html
        text = normal_html.decode("utf-8")
        headers = {"content-type": "text/html"}
        url = "https://mp.weixin.qq.com/s/xyz789"

    class _FakeClient:
        def __init__(self, *a: Any, **kw: Any) -> None:
            pass
        def __enter__(self):
            return self
        def __exit__(self, *exc):
            return False
        def get(self, url: str):
            return _FakeResp()

    monkeypatch.setattr("shared.text_loader.httpx.Client", _FakeClient)

    doc = load_url("https://mp.weixin.qq.com/s/xyz789")
    assert "parser" not in doc.meta  # 没走微信专门路径
    assert "普通正文" in doc.content


def test_load_url_wechat_title_by_class(monkeypatch: pytest.MonkeyPatch) -> None:
    """微信标题 h1 只有 class 无 id → 也能抽到标题（不兜底 URL）。"""
    wx_html = (
        '<html><body>'
        '<h1 class="rich_media_title">纯 class 标题</h1>'
        '<div id="js_content"><p>正文。</p></div>'
        '</body></html>'
    ).encode("utf-8")

    class _FakeResp:
        status_code = 200
        content = wx_html
        text = wx_html.decode("utf-8")
        headers = {"content-type": "text/html"}
        url = "https://mp.weixin.qq.com/s/class-only"

    class _FakeClient:
        def __init__(self, *a: Any, **kw: Any) -> None:
            pass
        def __enter__(self):
            return self
        def __exit__(self, *exc):
            return False
        def get(self, url: str):
            return _FakeResp()

    monkeypatch.setattr("shared.text_loader.httpx.Client", _FakeClient)

    doc = load_url("https://mp.weixin.qq.com/s/class-only")
    assert doc.title == "纯 class 标题"
    assert doc.meta["parser"] == "wechat"


def test_load_url_wechat_title_nested_span(monkeypatch: pytest.MonkeyPatch) -> None:
    """微信标题文字在 span 子元素里 → text_content() 能拿到完整标题。"""
    wx_html = (
        '<html><body>'
        '<h1 class="rich_media_title " id="activity-name">'
        '<span class="js_title_inner">嵌套标题内容</span></h1>'
        '<div id="js_content"><p>正文。</p></div>'
        '</body></html>'
    ).encode("utf-8")

    class _FakeResp:
        status_code = 200
        content = wx_html
        text = wx_html.decode("utf-8")
        headers = {"content-type": "text/html"}
        url = "https://mp.weixin.qq.com/s/nested-span"

    class _FakeClient:
        def __init__(self, *a: Any, **kw: Any) -> None:
            pass
        def __enter__(self):
            return self
        def __exit__(self, *exc):
            return False
        def get(self, url: str):
            return _FakeResp()

    monkeypatch.setattr("shared.text_loader.httpx.Client", _FakeClient)

    doc = load_url("https://mp.weixin.qq.com/s/nested-span")
    assert doc.title == "嵌套标题内容"
    assert doc.meta["parser"] == "wechat"


# ── handle_text_task 集成（mock load_auto）─────────────────


def _wait_terminal(store: TaskStore, task_id: str, timeout: float = 5.0):
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        rec = store.get(task_id)
        if rec is not None and rec.status in TERMINAL_STATUS_VALUES:
            return rec
        time.sleep(0.05)
    return store.get(task_id)


def test_handle_text_task_happy(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    # 改写 WORKSPACES_DATA_DIR，让产物落到 tmp 下
    monkeypatch.setattr("shared.config.WORKSPACES_DATA_DIR", tmp_path)
    # 注意：pipeline_tasks 是 from shared.config import get_workspace_text_dir
    # 而 get_workspace_text_dir 内部读 WORKSPACES_DATA_DIR 模块级常量，monkeypatch 有效。

    fake_doc = TextDocument(
        title="Fake Title",
        content="Fake body content " * 10,
        source_type="url",
        source="https://example.com/post",
        char_count=200,
        meta={"http_status": 200},
    )
    monkeypatch.setattr(
        "backend.app.services.pipeline_tasks.load_auto",
        lambda source, source_type=None: fake_doc,
    )
    # 让 _summarize_text 走"未提供 api_key"分支，返回空串
    fake_settings = MagicMock()
    fake_settings.openai_api_key = ""
    fake_settings.text_model = ""
    monkeypatch.setattr(
        "backend.app.services.pipeline_tasks.load_settings",
        lambda: fake_settings,
    )

    store = TaskStore(path=tmp_path / "tasks.json")
    runner = TaskRunner(store, max_workers=1)
    runner.register("text", handle_text_task)

    rec = runner.create_task("proj-text", "text", {"source": "https://example.com/post"})
    final = _wait_terminal(store, rec.task_id)

    assert final is not None
    assert final.status == TaskStatus.SUCCESS.value, final.error
    result = final.result
    assert result["title"] == "Fake Title"
    assert result["char_count"] == 200
    assert result["summary"] == {"abstract": "", "key_points": [], "golden_quotes": []}
    md_path = Path(result["markdown_path"])
    json_path = Path(result["json_path"])
    assert md_path.is_file() and json_path.is_file()
    assert "Fake Title" in md_path.read_text(encoding="utf-8")


def test_handle_text_task_missing_source(tmp_path: Path) -> None:
    store = TaskStore(path=tmp_path / "tasks.json")
    runner = TaskRunner(store, max_workers=1)
    runner.register("text", handle_text_task)

    rec = runner.create_task("proj-text", "text", {})  # 缺 source
    final = _wait_terminal(store, rec.task_id)
    assert final is not None
    assert final.status == TaskStatus.FAILED.value
    assert "source" in final.error.lower()


# ── _parse_structured_summary 单测 ──────────────────────────

SAMPLE_CONTENT = (
    "人工智能技术在过去十年取得了飞速发展。"
    "深度学习模型在图像识别、自然语言处理等领域表现卓越。"
    "然而，大模型的训练成本居高不下，成为学术界的一大挑战。"
    "未来，轻量化模型和绿色AI将成为重要研究方向。"
)


def _parse(raw: str) -> dict[str, Any]:
    """便捷包装，避免每次都传 content 和 log。"""
    log_lines: list[str] = []
    result = _parse_structured_summary(raw, SAMPLE_CONTENT, log_lines.append)
    result["_log"] = log_lines
    return result


class TestParseStructuredSummaryGoldenQuotes:
    """金句子串校验：必须是原文精确子串，否则丢弃。"""

    def test_valid_quote_gets_position(self) -> None:
        quote = "深度学习模型在图像识别、自然语言处理等领域表现卓越"
        raw = json.dumps({"golden_quotes": [{"quote_text": quote}]}, ensure_ascii=False)
        r = _parse(raw)
        assert len(r["golden_quotes"]) == 1
        q = r["golden_quotes"][0]
        assert q["quote_text"] == quote
        assert q["char_start"] >= 0
        assert q["char_end"] > q["char_start"]
        assert "para_index" in q

    def test_invalid_quote_discarded(self) -> None:
        raw = json.dumps({"golden_quotes": [{"quote_text": "这段文字根本不在原文里"}]}, ensure_ascii=False)
        r = _parse(raw)
        assert r["golden_quotes"] == []
        assert any("金句子串校验失败" in l for l in r["_log"])

    def test_empty_quote_skipped(self) -> None:
        raw = json.dumps({"golden_quotes": [{"quote_text": ""}]}, ensure_ascii=False)
        r = _parse(raw)
        assert r["golden_quotes"] == []

    def test_non_dict_entry_skipped(self) -> None:
        raw = json.dumps({"golden_quotes": ["not a dict"]}, ensure_ascii=False)
        r = _parse(raw)
        assert r["golden_quotes"] == []


class TestParseStructuredSummaryKeyPoints:
    """要点校验：source_excerpt 找不到时丢弃，找不到 excerpt 的保留为纯文本要点。"""

    def test_anchored_excerpt_gets_position(self) -> None:
        excerpt = "大模型的训练成本居高不下"
        raw = json.dumps(
            {"key_points": [{"text": "训练成本高", "source_excerpt": excerpt}]},
            ensure_ascii=False,
        )
        r = _parse(raw)
        assert len(r["key_points"]) == 1
        kp = r["key_points"][0]
        assert kp["text"] == "训练成本高"
        assert kp["source_excerpt"] == excerpt
        assert kp["char_start"] >= 0
        assert kp["char_end"] > kp["char_start"]
        assert "para_index" in kp

    def test_unmatched_excerpt_discarded(self) -> None:
        """source_excerpt 在原文中找不到 → 丢弃该要点。"""
        raw = json.dumps(
            {"key_points": [{"text": "某个要点", "source_excerpt": "这段文字不在原文里"}]},
            ensure_ascii=False,
        )
        r = _parse(raw)
        assert r["key_points"] == []
        assert any("source_excerpt 未在原文找到" in l for l in r["_log"])

    def test_no_excerpt_kept_as_plain_point(self) -> None:
        """无 source_excerpt 的要点保留，不丢弃。"""
        raw = json.dumps(
            {"key_points": [{"text": "纯文本要点，无摘录"}]},
            ensure_ascii=False,
        )
        r = _parse(raw)
        assert len(r["key_points"]) == 1
        kp = r["key_points"][0]
        assert kp["text"] == "纯文本要点，无摘录"
        assert "source_excerpt" not in kp
        assert "char_start" not in kp

    def test_mixed_valid_invalid_excerpts(self) -> None:
        """有效和无效 source_excerpt 混合时，仅保留有效的。"""
        raw = json.dumps(
            {
                "key_points": [
                    {"text": "有效要点", "source_excerpt": "轻量化模型和绿色AI"},
                    {"text": "无效要点", "source_excerpt": "编造的不存在原文"},
                    {"text": "无摘录要点"},
                ]
            },
            ensure_ascii=False,
        )
        r = _parse(raw)
        assert len(r["key_points"]) == 2
        texts = {kp["text"] for kp in r["key_points"]}
        assert texts == {"有效要点", "无摘录要点"}

    def test_empty_text_skipped(self) -> None:
        raw = json.dumps({"key_points": [{"text": ""}]}, ensure_ascii=False)
        r = _parse(raw)
        assert r["key_points"] == []

    def test_non_dict_entry_skipped(self) -> None:
        raw = json.dumps({"key_points": ["not a dict"]}, ensure_ascii=False)
        r = _parse(raw)
        assert r["key_points"] == []


# ── T1.2: _split_paragraphs 单测 ─────────────────────────────


class TestSplitParagraphs:
    def test_basic_split(self) -> None:
        text = "第一段。\n\n第二段，继续。\n\n第三段结尾。"
        result = _split_paragraphs(text)
        assert result == ["第一段。", "第二段，继续。", "第三段结尾。"]

    def test_multiple_newlines(self) -> None:
        text = "A\n\n\n\nB\n\n\nC"
        result = _split_paragraphs(text)
        assert result == ["A", "B", "C"]

    def test_trims_whitespace(self) -> None:
        text = "  A  \n\n  B  "
        result = _split_paragraphs(text)
        assert result == ["A", "B"]

    def test_filters_empty_segments(self) -> None:
        text = "\n\nA\n\n\n\nB\n\n"
        result = _split_paragraphs(text)
        assert result == ["A", "B"]

    def test_single_paragraph(self) -> None:
        text = "只有一段。"
        result = _split_paragraphs(text)
        assert result == ["只有一段。"]

    def test_empty_string(self) -> None:
        result = _split_paragraphs("")
        assert result == []


class TestParseStructuredSummaryEdgeCases:
    """边界情况：代码块包裹、无效 JSON、空输入。"""

    def test_json_wrapped_in_code_fence(self) -> None:
        excerpt = "人工智能技术在过去十年取得了飞速发展"
        payload = json.dumps(
            {"abstract": "摘要内容", "key_points": [{"text": "要点", "source_excerpt": excerpt}]},
            ensure_ascii=False,
        )
        raw = f"```json\n{payload}\n```"
        r = _parse(raw)
        assert r["abstract"] == "摘要内容"
        assert len(r["key_points"]) == 1

    def test_bare_code_fence_without_lang(self) -> None:
        excerpt = "深度学习模型"
        payload = json.dumps(
            {"key_points": [{"text": "要点", "source_excerpt": excerpt}]},
            ensure_ascii=False,
        )
        raw = f"```\n{payload}\n```"
        r = _parse(raw)
        assert len(r["key_points"]) == 1

    def test_invalid_json_falls_back_to_raw(self) -> None:
        r = _parse("这不是合法的 JSON 字符串")
        # abstract 回退到原始字符串
        assert r["abstract"] == "这不是合法的 JSON 字符串"
        assert r["key_points"] == []
        assert r["golden_quotes"] == []

    def test_truncated_json_with_braces(self) -> None:
        excerpt = "绿色AI"
        payload = json.dumps(
            {"key_points": [{"text": "要点", "source_excerpt": excerpt}]},
            ensure_ascii=False,
        )
        # 前后加废话，模拟 LLM 输出
        raw = f"好的，这是摘要：{payload}希望能帮到你。"
        r = _parse(raw)
        assert len(r["key_points"]) == 1

    def test_empty_string(self) -> None:
        r = _parse("")
        assert r["abstract"] == ""
        assert r["key_points"] == []
        assert r["golden_quotes"] == []

    def test_empty_object(self) -> None:
        r = _parse("{}")
        # parsed {} is truthy, abstract fallback reads raw "{}"
        assert r["abstract"] == "{}"
        assert r["key_points"] == []
        assert r["golden_quotes"] == []
