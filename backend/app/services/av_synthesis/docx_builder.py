"""R20: Word 导出 — python-docx 构建。

python-docx 已在 requirements.txt，无需额外安装。
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from fastapi.responses import StreamingResponse


def _add_image(doc: Any, ws_root: Path, rel_path: str, width_inches: float = 5.5) -> bool:
    """向文档插入图片。返回是否成功。"""
    if not rel_path:
        return False
    full = ws_root / rel_path.lstrip("./")
    if not full.exists():
        return False
    try:
        from docx.shared import Inches
        doc.add_picture(str(full), width=Inches(width_inches))
        return True
    except Exception:
        return False


def build_docx(notes: Any, ws_root: Path) -> StreamingResponse:
    """生成 .docx 并返回 StreamingResponse。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体（中文友好）
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    # ── 标题 ──────────────────────────────────────
    heading = doc.add_heading(notes.title or "综合笔记", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── 元信息 ────────────────────────────────────
    meta_parts = []
    if notes.platform:
        meta_parts.append(notes.platform)
    if notes.author:
        meta_parts.append(notes.author)
    if notes.duration_display:
        meta_parts.append(notes.duration_display)
    if notes.date_added:
        meta_parts.append(notes.date_added)
    if meta_parts:
        p = doc.add_paragraph(" · ".join(meta_parts))
        p.runs[0].font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
        p.runs[0].font.size = Pt(10)

    # ── 封面图 ────────────────────────────────────
    if notes.cover_path:
        _add_image(doc, ws_root, notes.cover_path, width_inches=6.0)

    # ── 全局摘要 ──────────────────────────────────
    doc.add_heading("全局摘要", level=1)
    if notes.summary:
        doc.add_paragraph(notes.summary)

    # ── 关键帧画廊 ────────────────────────────────
    if notes.gallery_rows:
        doc.add_heading("关键帧画廊", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "时刻"
        hdr[1].text = "画面"
        hdr[2].text = "场景描述"
        for row in notes.gallery_rows:
            cells = table.add_row().cells
            cells[0].text = row.timestamp_display
            # 图片插入到单元格
            if row.image_path:
                full = ws_root / row.image_path.lstrip("./")
                if full.exists():
                    try:
                        from docx.shared import Inches
                        cells[1].paragraphs[0].clear()
                        run = cells[1].paragraphs[0].add_run()
                        run.add_picture(str(full), width=Inches(1.2))
                    except Exception:
                        cells[1].text = "-"
                else:
                    cells[1].text = "-"
            else:
                cells[1].text = "-"
            cells[2].text = row.scene_description

    # ── 章节正文 ──────────────────────────────────
    if notes.chapters:
        doc.add_heading("章节正文", level=1)
        for i, ch in enumerate(notes.chapters, 1):
            doc.add_heading(f"{i}. {ch.title}（{ch.time_range}）", level=2)
            if ch.frame_path:
                _add_image(doc, ws_root, ch.frame_path, width_inches=5.0)
            if ch.transcript_excerpt:
                p = doc.add_paragraph()
                p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else doc.styles["Normal"]
                p.add_run(ch.transcript_excerpt).font.size = Pt(10)
            if ch.highlights:
                p = doc.add_paragraph()
                run = p.add_run(f"重点：{ch.highlights}")
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x4D, 0x7E)

    # ── 字幕原文 ──────────────────────────────────
    if notes.full_transcript:
        doc.add_heading("字幕原文", level=1)
        # 分段添加避免超长段落
        for line in notes.full_transcript.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    # ── 最终综合 ──────────────────────────────────
    if notes.final_synthesis:
        doc.add_heading("最终综合", level=1)
        doc.add_paragraph(notes.final_synthesis)

    # ── 输出 ──────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from urllib.parse import quote
    safe_title = (notes.title or "笔记").replace("/", "_").replace("\\", "_")[:50]
    filename = f"{safe_title}.docx"
    filename_star = f"UTF-8''{quote(filename)}"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*= {filename_star}",
        },
    )
